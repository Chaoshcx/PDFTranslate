import asyncio
import aiohttp
import requests
from docx import Document
import os
import json
import re
import time
import pdfplumber
import argparse
from pathlib import Path
from dotenv import load_dotenv
from tqdm import tqdm
import logging
import sys
import filelock  # 导入文件锁模块

# 配置日志
logging.basicConfig(
    level=logging.DEBUG,  # 将日志级别改为DEBUG，以记录更详细的信息
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("translation_debug.log",
                            encoding='utf-8'),  # 添加encoding参数确保中文正确显示
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("translator")

# 添加一个专门用于段落对照的日志处理器
paragraph_logger = logging.getLogger("paragraph_comparison")
paragraph_logger.setLevel(logging.DEBUG)
paragraph_handler = logging.FileHandler(
    "paragraph_comparison.log", encoding='utf-8')
paragraph_handler.setFormatter(logging.Formatter(
    '%(asctime)s - %(levelname)s - %(message)s'))
paragraph_logger.addHandler(paragraph_handler)
paragraph_logger.propagate = False  # 防止日志消息传递到根日志记录器

# 尝试加载环境变量
try:
    load_dotenv()
except ImportError:
    logger.warning("dotenv 模块未安装，跳过环境变量加载")

# 默认配置 - 移除个人路径和API密钥
DEFAULT_CONFIG = {
    "pdf_path": "",
    "output_dir": "./outputs",
    "progress_file": "./outputs/translation_progress.json",
    "api_key": os.getenv("API_KEY", ""),  # 默认为空，由用户填写
    "api_url": "https://ark.cn-beijing.volces.com/api/v3/chat/completions",
    "api_model": "deepseek-v3-241226",
    "batch_size": 3,
    "max_concurrent_requests": 3,
    "temperature": 0.3,
    "retries": 3,
    "retry_delay": 2,
    "paragraphs_per_page": 10
}

# 尝试导入自定义配置文件
try:
    import config
    # 更新默认配置
    for key in DEFAULT_CONFIG:
        if hasattr(config, key.upper()):
            DEFAULT_CONFIG[key] = getattr(config, key.upper())
    logger.info("已加载自定义配置")
except ImportError:
    logger.info("未找到自定义配置文件，使用默认配置")


def show_progress_bar(current, total, width=50):
    """显示漂亮的进度条"""
    progress = min(1.0, current / total)
    filled_width = int(width * progress)
    bar = '█' * filled_width + '░' * (width - filled_width)
    percent = progress * 100
    return f"\r进度: |{bar}| {percent:.1f}% 完成 ({current}/{total})"


def extract_text_from_pdf(pdf_path):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = ""
            total_pages = len(pdf.pages)
            logger.info(f"开始从PDF提取文本，共 {total_pages} 页")

            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                text += page_text + "\n"

                # 显示进度
                if i % 5 == 0 or i == total_pages - 1:
                    print(show_progress_bar(i + 1, total_pages), end='')

            print()  # 换行
            logger.info(f"成功从PDF提取文本，长度: {len(text)} 字符")
            return text
    except Exception as e:
        logger.error(f"PDF提取失败: {e}")
        return ""


def get_pdf_preview(pdf_path, max_pages=3, max_paragraphs=5):
    """从PDF中提取预览文本，返回前几页的几个段落"""
    try:
        preview_data = {
            "total_pages": 0,
            "previews": []
        }

        with pdfplumber.open(pdf_path) as pdf:
            preview_data["total_pages"] = len(pdf.pages)
            pages_to_preview = min(max_pages, len(pdf.pages))

            for i in range(pages_to_preview):
                page = pdf.pages[i]
                page_text = page.extract_text() or ""

                # 分割段落
                paragraphs = [p for p in page_text.split('\n\n') if p.strip()]

                # 取前几个段落
                preview_paragraphs = paragraphs[:max_paragraphs]

                preview_data["previews"].append({
                    "page": i + 1,
                    "paragraphs": preview_paragraphs
                })

            return preview_data
    except Exception as e:
        logger.error(f"PDF预览提取失败: {e}")
        return {"total_pages": 0, "previews": []}


def save_text_to_docx(text, docx_path):
    try:
        doc = Document()
        paragraphs = text.split("\n")
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para)
        doc.save(docx_path)
        logger.info(f"成功保存临时DOCX文件: {docx_path}")
    except Exception as e:
        logger.error(f"保存DOCX失败: {e}")


def extract_text_from_word(word_path):
    try:
        doc = Document(word_path)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        logger.info(f"成功从Word提取文本，长度: {len(text)} 字符")
        return text
    except Exception as e:
        logger.error(f"Word提取失败: {e}")
        return ""


async def translate_with_deepseek_async(paragraphs, api_key, api_url, api_model, temperature=0.3, retries=3, delay=2):
    """异步版本的deepseek翻译函数"""
    # 改进系统提示以获得更好的翻译并保持段落结构
    system_prompt = """你是一个专业的翻译助手。请将以下英文文本翻译成中文。
    要求：
    1. 保持专业术语的准确性
    2. 确保翻译通顺易读
    3. 严格保持段落结构，每个段落都必须对应翻译
    4. 确保翻译准确无误，不要遗漏内容
    5. 如有多个段落，请在翻译时保持段落之间的分隔，即每个段落一行
    6. 不要合并或拆分段落，确保输入和输出的段落数量相同
    7. 每个翻译段落必须单独成行，使用换行符分隔
    
    重要：请只输出翻译结果，不要输出任何其他内容，如注释、说明、分析或翻译过程等。
    """

    # 记录原始段落数量和内容
    logger.info(f"待翻译的段落数: {len(paragraphs)}")
    for i, para in enumerate(paragraphs):
        logger.debug(f"待翻译-段落{i+1}: {para[:50]}...")

    # 对每个段落添加标记，以便在响应中更容易识别
    marked_paragraphs = []
    for i, para in enumerate(paragraphs):
        marked_paragraphs.append(f"段落{i+1}: {para}")

    paragraph_text = "\n\n".join(marked_paragraphs)

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"请将以下文本翻译成中文，直接给出翻译结果，不要添加任何解释、注释或备注。请确保每个段落都有对应的翻译，并保持段落顺序。每个段落必须单独成行，段落之间使用换行符分隔:\n\n{paragraph_text}"}
    ]

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }

    data = {
        "model": api_model,
        "messages": messages,
        "temperature": temperature
    }

    for attempt in range(retries):
        try:
            logger.info(f"正在进行第 {attempt + 1} 次翻译尝试...")
            async with aiohttp.ClientSession() as session:
                async with session.post(api_url, json=data, headers=headers) as response:
                    if response.status != 200:
                        error_text = await response.text()
                        logger.error(
                            f"API响应错误 (HTTP {response.status}): {error_text}")
                        raise Exception(f"API响应错误: {response.status}")

                    response_json = await response.json()
                    if "choices" not in response_json or not response_json["choices"]:
                        logger.error(f"API响应格式错误: {response_json}")
                        raise Exception("API响应格式错误")

                    translated_text = response_json["choices"][0]["message"]["content"]
                    logger.debug(f"API原始返回结果: {translated_text[:200]}...")

                    # 处理返回的文本，尝试恢复段落结构
                    processed_text = translated_text

                    # 删除可能的"段落X:"前缀
                    logger.debug("正在处理译文，删除段落标记...")
                    processed_text = re.sub(
                        r'^段落\d+[:：]\s*', '', processed_text, flags=re.MULTILINE)

                    # 记录处理后的段落数
                    raw_paragraphs = processed_text.split('\n')
                    non_empty_paragraphs = [
                        p for p in raw_paragraphs if p.strip()]
                    logger.info(
                        f"API返回段落处理: 总行数={len(raw_paragraphs)}, 非空行数={len(non_empty_paragraphs)}")

                    # 确保段落之间有换行符
                    if len(paragraphs) > 1 and "\n\n" not in processed_text and "\n" not in processed_text:
                        logger.warning("翻译结果未包含段落分隔符，尝试使用原始段落数量分割...")

                        # 如果没有换行符，尝试按句子分割并重组为与原段落数相匹配的结构
                        sentences = re.split(
                            r'(?<=[。！？.!?])\s*', processed_text)
                        sentences = [s for s in sentences if s.strip()]
                        logger.debug(f"分割得到 {len(sentences)} 个句子")

                        if len(sentences) >= len(paragraphs):
                            # 估算每个段落的句子数
                            sentences_per_para = max(
                                1, len(sentences) // len(paragraphs))
                            logger.debug(f"估算每个段落包含 {sentences_per_para} 个句子")

                            new_paragraphs = []
                            for i in range(0, len(sentences), sentences_per_para):
                                para_sentences = sentences[i:i +
                                                           sentences_per_para]
                                new_paragraphs.append("".join(para_sentences))

                            # 确保段落数量不超过原文
                            new_paragraphs = new_paragraphs[:len(paragraphs)]
                            processed_text = "\n\n".join(new_paragraphs)
                            logger.info(
                                f"重新格式化翻译结果为 {len(new_paragraphs)} 个段落")

                            # 记录重组后的段落
                            for i, para in enumerate(new_paragraphs):
                                logger.debug(f"重组后-段落{i+1}: {para[:50]}...")
                    else:
                        # 规范化段落分隔符
                        processed_text = re.sub(r'\n+', '\n\n', processed_text)
                        logger.debug("使用规范化的段落分隔符处理文本")

                    # 最终检查段落数量
                    final_paragraphs = [
                        p for p in processed_text.split('\n\n') if p.strip()]
                    logger.info(
                        f"最终段落数: {len(final_paragraphs)}, 原始段落数: {len(paragraphs)}")

                    for i, para in enumerate(final_paragraphs):
                        logger.debug(f"最终-段落{i+1}: {para[:50]}...")

                    preview = processed_text[:100] + "..." if len(
                        processed_text) > 100 else processed_text
                    logger.info(f"翻译成功，结果预览: {preview}")
                    return processed_text
        except Exception as e:
            logger.error(f"翻译请求失败 (尝试 {attempt + 1}): {e}")
            if attempt < retries - 1:
                logger.info(f"等待 {delay} 秒后重试...")
                await asyncio.sleep(delay)
            else:
                logger.error("所有重试都失败，跳过当前段落")
                return ""


def translate_with_deepseek(paragraphs, api_key, api_url, api_model, temperature=0.3, retries=3, delay=2):
    """同步版本的deepseek翻译函数（保留以便兼容）"""
    # 改进系统提示以获得更好的翻译
    system_prompt = """你是一个专业的翻译助手。请将以下英文文本翻译成中文。
    要求：
    1. 保持专业术语的准确性
    2. 确保翻译通顺易读
    3. 严格保持段落结构，每个段落都必须对应翻译
    4. 确保翻译准确无误，不要遗漏内容
    
    重要：请只输出翻译结果，不要输出任何其他内容，如注释、说明、分析或翻译过程等。
    """

    # 对每个段落添加标记，以便在响应中更容易识别
    marked_paragraphs = []
    for i, para in enumerate(paragraphs):
        marked_paragraphs.append(f"段落{i+1}: {para}")

    paragraph_text = "\n\n".join(marked_paragraphs)

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"请将以下文本翻译成中文，直接给出翻译结果，不要添加任何解释、注释或备注。请确保每个段落都有对应的翻译，并保持段落顺序:\n\n{paragraph_text}"}
    ]

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }

    data = {
        "model": api_model,
        "messages": messages,
        "temperature": temperature
    }

    for attempt in range(retries):
        try:
            logger.info(f"正在进行第 {attempt + 1} 次翻译尝试...")
            response = requests.post(api_url, json=data, headers=headers)
            response.raise_for_status()
            response_json = response.json()
            translated_text = response_json["choices"][0]["message"]["content"]

            # 处理返回的文本，尝试恢复段落结构
            processed_text = translated_text

            # 删除可能的"段落X:"前缀
            processed_text = re.sub(
                r'^段落\d+[:：]\s*', '', processed_text, flags=re.MULTILINE)

            # 确保段落之间有换行符
            if len(paragraphs) > 1 and "\n\n" not in processed_text and "\n" not in processed_text:
                logger.warning("翻译结果未包含段落分隔符，尝试使用原始段落数量分割...")

                # 如果没有换行符，尝试按句子分割并重组为与原段落数相匹配的结构
                sentences = re.split(r'(?<=[。！？.!?])\s*', processed_text)
                sentences = [s for s in sentences if s.strip()]

                if len(sentences) >= len(paragraphs):
                    # 估算每个段落的句子数
                    sentences_per_para = max(
                        1, len(sentences) // len(paragraphs))
                    new_paragraphs = []

                    for i in range(0, len(sentences), sentences_per_para):
                        para_sentences = sentences[i:i+sentences_per_para]
                        new_paragraphs.append("".join(para_sentences))

                    # 确保段落数量不超过原文
                    new_paragraphs = new_paragraphs[:len(paragraphs)]
                    processed_text = "\n\n".join(new_paragraphs)
                    logger.info(f"重新格式化翻译结果为 {len(new_paragraphs)} 个段落")

            preview = processed_text[:100] + \
                "..." if len(processed_text) > 100 else processed_text
            logger.info(f"翻译成功，结果预览: {preview}")
            return processed_text
        except requests.exceptions.RequestException as e:
            logger.error(f"翻译请求失败 (尝试 {attempt + 1}): {e}")
            if attempt < retries - 1:
                logger.info(f"等待 {delay} 秒后重试...")
                time.sleep(delay)
            else:
                logger.error("所有重试都失败，跳过当前段落")
                return ""


def save_to_word(translated_texts, output_path):
    """保存翻译结果到Word文档"""
    try:
        doc = Document()
        for text in translated_texts:
            # 分割并处理每个段落
            paragraphs = text.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
        doc.save(output_path)
        logger.info(f"成功保存翻译结果到: {output_path}")
        logger.info(f"文档包含 {len(translated_texts)} 个翻译段落")
    except Exception as e:
        logger.error(f"保存翻译结果失败: {e}")


def save_to_word_with_comparison(original_paragraphs, translated_texts, output_path):
    """保存原文与翻译对照的Word文档"""
    try:
        doc = Document()

        # 添加标题
        doc.add_heading('翻译对照文档', 0)

        # 预处理原始段落，过滤掉空段落
        valid_original_paragraphs = [p for p in original_paragraphs if p.strip()]
        original_count = len(valid_original_paragraphs)
        
        # 处理翻译后的文本批次
        all_translated_paragraphs = []
        
        # 记录原文段落数量
        paragraph_logger.info(f"原文有效段落数: {original_count}")
        
        # 处理所有翻译批次的结果，确保结果中不包含空段落
        for batch_idx, batch_text in enumerate(translated_texts):
            # 分割批次翻译结果为段落并过滤空段落
            batch_paragraphs = [p.strip() for p in batch_text.split('\n') if p.strip()]
            paragraph_logger.debug(f"批次{batch_idx+1}分割出 {len(batch_paragraphs)} 个段落")
            all_translated_paragraphs.extend(batch_paragraphs)
        
        # 记录处理后的译文段落数量
        translated_count = len(all_translated_paragraphs)
        paragraph_logger.info(f"处理后的译文段落数: {translated_count}")
        
        # 确定实际可以构建的对照段落数量 - 使用最小值以避免索引越界
        usable_paragraphs = min(original_count, translated_count)
        paragraph_logger.info(f"将创建 {usable_paragraphs} 对翻译对照")
        
        # 写入段落对照详细日志
        comparison_log_path = output_path.replace('.docx', '_段落对照.log')
        with open(comparison_log_path, 'w', encoding='utf-8') as log_file:
            log_file.write("===== 段落对照详细日志 =====\n\n")
            log_file.write(f"原文段落数: {original_count}, 译文段落数: {translated_count}\n")
            log_file.write(f"实际对照段落数: {usable_paragraphs}\n\n")
            
            # 记录每对段落的对照情况
            for i in range(usable_paragraphs):
                log_file.write(f"------ 段落 {i+1} ------\n")
                log_file.write(f"原文: {valid_original_paragraphs[i]}\n\n")
                log_file.write(f"译文: {all_translated_paragraphs[i]}\n\n")
            
            # 记录未匹配的原文段落
            if original_count > usable_paragraphs:
                log_file.write(f"\n===== 未翻译的 {original_count - usable_paragraphs} 个原文段落 =====\n\n")
                for i in range(usable_paragraphs, original_count):
                    log_file.write(f"未翻译段落 {i+1}: {valid_original_paragraphs[i]}\n\n")
            
            # 记录多余的译文段落
            if translated_count > usable_paragraphs:
                log_file.write(f"\n===== 多余的 {translated_count - usable_paragraphs} 个译文段落 =====\n\n")
                for i in range(usable_paragraphs, translated_count):
                    log_file.write(f"多余译文段落 {i+1}: {all_translated_paragraphs[i]}\n\n")
        
        # 创建Word文档内容 - 按匹配的段落创建对照
        for i in range(usable_paragraphs):
            # 添加原文
            doc.add_heading(f'段落 {i+1} - 原文', level=2)
            doc.add_paragraph(valid_original_paragraphs[i])
            
            # 添加译文
            doc.add_heading(f'段落 {i+1} - 译文', level=2)
            doc.add_paragraph(all_translated_paragraphs[i])
            
            # 添加分隔线（除了最后一段）
            if i < usable_paragraphs - 1:
                doc.add_paragraph('---')
        
        # 保存文档
        doc.save(output_path)
        paragraph_logger.info(f"成功保存翻译对照结果到: {output_path}")
        paragraph_logger.info(f"段落对照详细日志已保存至: {comparison_log_path}")
        logger.info(f"成功保存翻译对照结果到: {output_path}")
        logger.info(f"文档包含 {usable_paragraphs} 对原文和译文")
        
        return True
    except Exception as e:
        logger.error(f"保存翻译对照结果失败: {e}")
        raise


def split_text_into_batches(text, batch_size):
    """将文本分割成多个批次"""
    paragraphs = text.split("\n\n")
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    # 按照指定的批次大小分组
    batches = []
    for i in range(0, len(paragraphs), batch_size):
        batch = paragraphs[i:i+batch_size]
        batches.append(batch)

    return batches, paragraphs


def split_paragraphs_optimized(text, min_length=50):
    """优化的段落分割方法，避免段落过短"""
    # 使用双换行符分割基本段落
    paragraphs = text.split("\n\n")
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    # 合并过短的段落
    optimized_paragraphs = []
    current_paragraph = ""

    for p in paragraphs:
        if len(current_paragraph) < min_length:
            if current_paragraph:
                current_paragraph += "\n\n" + p
            else:
                current_paragraph = p
        else:
            optimized_paragraphs.append(current_paragraph)
            current_paragraph = p

    # 添加最后一个段落
    if current_paragraph:
        optimized_paragraphs.append(current_paragraph)

    return optimized_paragraphs


def split_paragraphs_by_sentences(text, sentences_per_paragraph=4):
    """根据句号分割文本，并按照指定规则形成段落

    规则：
    1. 以句号（.）作为句子的基本断点
    2. 如果句号后面紧跟换行符，则该句子单独成为一个段落
    3. 否则，按照每四句一个段落进行分组
    """
    logger.info("使用基于句号的段落分割逻辑...")

    # 预处理文本，规范化换行符
    text = text.replace('\r\n', '\n')

    # 定义英文句子结束的模式：句号、问号或感叹号，后跟空格或换行
    sentence_end_pattern = r'[.!?][\s\n]'

    # 查找所有可能的句子断点
    sentence_breaks = []
    for match in re.finditer(sentence_end_pattern, text):
        sentence_breaks.append(match.end() - 1)  # 减1是为了指向空格或换行符前的标点

    # 如果没找到句子断点，尝试使用换行符作为备选
    if not sentence_breaks:
        logger.warning("未检测到有效句子断点，尝试使用换行符分割...")
        lines = text.split('\n')
        filtered_lines = [line.strip() for line in lines if line.strip()]
        logger.info(f"使用换行符分割结果: {len(filtered_lines)} 行")
        return filtered_lines

    # 根据句子断点分割文本
    sentences = []
    prev_end = 0

    for end in sentence_breaks:
        # 从上一个断点到当前断点形成一个句子
        sentence = text[prev_end:end + 1].strip()
        if sentence:
            sentences.append(sentence)
        prev_end = end + 1

    # 添加最后一部分（如果有）
    if prev_end < len(text):
        last_sentence = text[prev_end:].strip()
        if last_sentence:
            sentences.append(last_sentence)

    logger.info(f"检测到 {len(sentences)} 个句子")

    # 形成段落
    paragraphs = []
    current_paragraph = []

    for sentence in sentences:
        current_paragraph.append(sentence)

        # 如果句子结尾有换行符，或者已经达到指定的句子数量，则形成一个段落
        if sentence.endswith('\n') or len(current_paragraph) >= sentences_per_paragraph:
            paragraph_text = ' '.join(current_paragraph).strip()
            if paragraph_text:
                paragraphs.append(paragraph_text)
            current_paragraph = []

    # 处理剩余的句子（如果有）
    if current_paragraph:
        paragraph_text = ' '.join(current_paragraph).strip()
        if paragraph_text:
            paragraphs.append(paragraph_text)

    logger.info(f"生成了 {len(paragraphs)} 个段落")

    # 过滤空段落并返回
    result = [p for p in paragraphs if p.strip()]
    return result


async def translate_batch_async(batch, api_key, api_url, api_model, semaphore, progress_data, progress_file):
    """异步翻译一个批次的段落"""
    async with semaphore:
        try:
            # 记录批次信息用于调试
            batch_info = f"批次索引: {progress_data.get('batch_index', 'N/A')}, 批次大小: {len(batch)}"
            logger.info(f"开始处理{batch_info}")

            # 记录原文内容，便于调试
            batch_idx = progress_data.get('batch_index', 'unknown')
            for i, para in enumerate(batch):
                logger.debug(f"[批次{batch_idx}-段落{i}] 原文: {para[:100]}...")

            # 执行翻译
            translated_text = await translate_with_deepseek_async(
                batch, api_key, api_url, api_model
            )

            # 记录译文并与原文对照
            if translated_text:
                # 分割译文为段落
                translated_paragraphs = []
                for p in translated_text.split('\n'):
                    if p.strip():
                        translated_paragraphs.append(p.strip())

                # 记录译文段落数与原文段落数的差异
                logger.info(
                    f"[批次{batch_idx}] 原文段落数: {len(batch)}, 译文段落数: {len(translated_paragraphs)}")

                # 记录译文内容与原文对照
                min_len = min(len(batch), len(translated_paragraphs))
                for i in range(min_len):
                    logger.debug(
                        f"[批次{batch_idx}-段落{i}] 原文: {batch[i][:50]}...")
                    logger.debug(
                        f"[批次{batch_idx}-段落{i}] 译文: {translated_paragraphs[i][:50]}...")

                # 记录多余段落
                if len(translated_paragraphs) > len(batch):
                    for i in range(len(batch), len(translated_paragraphs)):
                        logger.warning(
                            f"[批次{batch_idx}] 多余译文段落{i}: {translated_paragraphs[i][:50]}...")
                elif len(translated_paragraphs) < len(batch):
                    for i in range(len(translated_paragraphs), len(batch)):
                        logger.warning(
                            f"[批次{batch_idx}] 未翻译原文段落{i}: {batch[i][:50]}...")

            # 使用文件锁确保多任务共享进度数据的一致性
            lock_file = f"{progress_file}.lock"
            lock = filelock.FileLock(lock_file, timeout=30)  # 30秒超时

            try:
                with lock:
                    # 重新读取最新的进度数据
                    try:
                        with open(progress_file, 'r', encoding='utf-8') as f:
                            current_progress = json.load(f)
                    except (FileNotFoundError, json.JSONDecodeError):
                        # 如果文件不存在或格式错误，初始化进度数据
                        current_progress = {
                            'total': progress_data['total'],
                            'processed': [],
                            'batch_size': progress_data.get('batch_size', len(batch))
                        }

                    # 更新进度
                    batch_index = progress_data.get('batch_index', 0)
                    batch_size = progress_data.get('batch_size', len(batch))

                    # 计算要更新的索引
                    indices_to_add = []
                    for i in range(len(batch)):
                        index = batch_index * batch_size + i
                        if index < current_progress['total'] and index not in current_progress['processed']:
                            indices_to_add.append(index)

                    # 添加新处理的索引
                    current_progress['processed'].extend(indices_to_add)

                    # 保存进度数据
                    with open(progress_file, 'w', encoding='utf-8') as f:
                        json.dump(current_progress, f)

                    # 输出进度信息
                    completed = len(current_progress['processed'])
                    total = current_progress['total']
                    percentage = (completed / total) * 100 if total > 0 else 0
                    logger.info(
                        f"进度更新: {completed}/{total} ({percentage:.1f}%) - {batch_info}")
            except filelock.Timeout:
                logger.warning(f"无法获取进度文件锁，跳过进度更新 - {batch_info}")

            return translated_text
        except Exception as e:
            logger.error(
                f"批次处理失败: {e} - {batch_info if 'batch_info' in locals() else '未知批次'}")
            return ""


async def main_async(pdf_path, api_key, output_path, progress_file=None,
                     start_page=1, end_page=None, comparison_mode=False,
                     batch_size=3, max_concurrent_requests=3,
                     api_url=DEFAULT_CONFIG['api_url'],
                     api_model=DEFAULT_CONFIG['api_model'],
                     sentences_per_paragraph=4):
    """主异步翻译函数"""
    # 提取文本
    text = extract_text_from_pdf(pdf_path)
    if not text:
        logger.error("PDF文本提取失败，退出翻译")
        return False

    # 使用基于句号的段落分割
    paragraphs = split_paragraphs_by_sentences(text, sentences_per_paragraph)
    total_paragraphs = len(paragraphs)
    logger.info(f"共分割出 {total_paragraphs} 个段落")

    # 检查进度文件
    progress_data = {'total': total_paragraphs, 'processed': []}
    if progress_file and os.path.exists(progress_file):
        try:
            with open(progress_file, 'r', encoding='utf-8') as f:
                progress_data = json.load(f)
                logger.info(
                    f"已加载翻译进度，已处理 {len(progress_data['processed'])} 个段落")
        except:
            logger.warning("进度文件损坏或无法读取，创建新的进度文件")

    # 创建批次
    batches = []
    for i in range(0, total_paragraphs, batch_size):
        batch = paragraphs[i:i+batch_size]
        batches.append(batch)

    total_batches = len(batches)
    logger.info(f"将 {total_paragraphs} 个段落分为 {total_batches} 个批次进行处理")

    # 创建进度文件
    if progress_file:
        progress_data['total'] = total_paragraphs
        progress_data['batch_size'] = batch_size
        with open(progress_file, 'w', encoding='utf-8') as f:
            json.dump(progress_data, f)

    # 创建任务
    semaphore = asyncio.Semaphore(max_concurrent_requests)
    tasks = []

    # 创建一个映射，用于追踪每个批次对应的原始段落索引
    batch_to_original_indices = {}

    for i, batch in enumerate(batches):
        # 记录批次和原始段落的映射关系
        start_idx = i * batch_size
        end_idx = min(start_idx + len(batch), total_paragraphs)
        batch_to_original_indices[i] = list(range(start_idx, end_idx))

        # 检查这个批次是否需要处理
        skip_batch = True
        for j in range(len(batch)):
            idx = i * batch_size + j
            if idx < total_paragraphs and idx not in progress_data.get('processed', []):
                skip_batch = False
                break

        if skip_batch:
            logger.info(f"批次 {i+1}/{total_batches} 已处理，跳过")
            continue

        # 更新批次索引信息 - 为每个批次创建独立的进度数据副本
        batch_progress_data = progress_data.copy()
        batch_progress_data['batch_index'] = i

        # 创建翻译任务
        task = asyncio.create_task(
            translate_batch_async(
                batch, api_key, api_url, api_model,
                semaphore, batch_progress_data, progress_file
            )
        )
        tasks.append((i, task))  # 存储批次索引和任务

    # 执行所有任务
    translated_results = []
    original_order = []  # 保持原始顺序

    if tasks:
        logger.info(f"开始执行 {len(tasks)} 个翻译任务")
        # 按批次顺序收集结果
        for batch_idx, task in sorted(tasks, key=lambda x: x[0]):
            try:
                result = await task
                if result:  # 有效结果
                    translated_results.append(result)
                    # 记录这个批次对应的原始段落索引
                    if batch_idx in batch_to_original_indices:
                        original_order.extend(
                            batch_to_original_indices[batch_idx])
            except Exception as e:
                logger.error(f"批次 {batch_idx} 处理失败: {e}")

    # 保存结果
    if translated_results:
        if comparison_mode:
            # 在对照模式下，根据original_order重新组织原始段落
            ordered_paragraphs = []
            for idx in original_order:
                if 0 <= idx < len(paragraphs):
                    ordered_paragraphs.append(paragraphs[idx])

            logger.info(
                f"对照模式: 重新组织了 {len(ordered_paragraphs)} 个原始段落与 {len(translated_results)} 个翻译结果")
            save_to_word_with_comparison(
                ordered_paragraphs, translated_results, output_path)
        else:
            save_to_word(translated_results, output_path)

        logger.info(f"翻译完成！结果已保存至: {output_path}")
    else:
        logger.warning("没有生成翻译结果，请检查是否有错误发生")

    return True


def main():
    """命令行入口函数"""
    parser = argparse.ArgumentParser(description="PDF文档翻译工具")
    parser.add_argument("--pdf", help="PDF文件路径",
                        default=DEFAULT_CONFIG["pdf_path"])
    parser.add_argument("--api-key", help="API密钥",
                        default=DEFAULT_CONFIG["api_key"])
    parser.add_argument("--output", help="输出文件路径", default="")
    parser.add_argument("--comparison", help="是否生成对照文档", action="store_true")
    parser.add_argument("--batch", help="批处理大小", type=int,
                        default=DEFAULT_CONFIG["batch_size"])
    parser.add_argument("--concurrent", help="并发请求数", type=int,
                        default=DEFAULT_CONFIG["max_concurrent_requests"])
    parser.add_argument("--sentences", help="每个段落包含的句子数", type=int, default=4)

    args = parser.parse_args()

    # 设置输出路径
    if not args.output:
        pdf_name = os.path.basename(args.pdf).rsplit('.', 1)[0]
        output_dir = DEFAULT_CONFIG["output_dir"]
        os.makedirs(output_dir, exist_ok=True)
        args.output = os.path.join(output_dir, f"translated_{pdf_name}.docx")

    # 设置进度文件
    progress_file = os.path.join(
        DEFAULT_CONFIG["output_dir"], f"progress_{os.path.basename(args.pdf).rsplit('.', 1)[0]}.json")

    # 运行翻译
    asyncio.run(main_async(
        pdf_path=args.pdf,
        api_key=args.api_key,
        output_path=args.output,
        progress_file=progress_file,
        comparison_mode=args.comparison,
        batch_size=args.batch,
        max_concurrent_requests=args.concurrent,
        sentences_per_paragraph=args.sentences
    ))

    print(f"翻译完成！结果已保存到: {args.output}")


if __name__ == "__main__":
    main()
